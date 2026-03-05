"""
Build the CG-IDF v2 Word report.
Run with: python3 build_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x5C)
BLUE   = RGBColor(0x2E, 0x6D, 0xB4)
TEAL   = RGBColor(0x0D, 0x7C, 0x8C)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREY   = RGBColor(0x60, 0x60, 0x60)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xEE, 0xF4, 0xFB)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREEN  = RGBColor(0x1E, 0x8B, 0x4C)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        val = kwargs.get(edge, {})
        if val:
            tag = OxmlElement(f"w:{edge}")
            for k, v in val.items():
                tag.set(qn(f"w:{k}"), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color=NAVY, space_before=18, space_after=6):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run  = p.add_run(text)
    sizes = {1: 22, 2: 16, 3: 13}
    run.font.size  = Pt(sizes.get(level, 13))
    run.font.color.rgb = color
    run.font.bold  = True
    run.font.name  = "Calibri"
    return p

def add_body(text, space_after=4, italic=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.name  = "Calibri"
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, bold_prefix=None):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.bold  = True
        run.font.size  = Pt(11)
        run.font.name  = "Calibri"
        rest = p.add_run(text)
        rest.font.size = Pt(11)
        rest.font.name = "Calibri"
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p

def add_code_block(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F4F8")
    pPr.append(shd)
    return p

def add_table(headers, rows, header_color="1A3A5C", alt_color="EEF4FB"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(10)
        run.font.name  = "Calibri"
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)

    # Data rows
    for ri, row in enumerate(rows):
        bg = alt_color if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p   = cell.paragraphs[0]
            if isinstance(val, tuple):          # (text, bold)
                run = p.add_run(val[0])
                run.font.bold = val[1]
            else:
                run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

    doc.add_paragraph()   # spacing after table
    return t

def add_divider():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DB4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)

def add_callout(text, bg_hex="EEF4FB", border_hex="2E6DB4"):
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.2)
    p.paragraph_format.space_before  = Pt(6)
    p.paragraph_format.space_after   = Pt(10)
    run = p.add_run(text)
    run.font.size   = Pt(10.5)
    run.font.italic = True
    run.font.name   = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg_hex)
    pPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CG-IDF v2")
run.font.size  = Pt(36)
run.font.bold  = True
run.font.color.rgb = NAVY
run.font.name  = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Incentive Design Forensic Engine")
run.font.size  = Pt(18)
run.font.color.rgb = BLUE
run.font.name  = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(60)
run = p.add_run("Version 2.0.0  |  Application Architecture & Output Report")
run.font.size  = Pt(12)
run.font.color.rgb = GREY
run.font.name  = "Calibri"

add_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("March 2026")
run.font.size  = Pt(11)
run.font.color.rgb = GREY
run.font.name  = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS CG-IDF v2?
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1. What Is CG-IDF v2?", level=1)
add_divider()
add_body(
    "CG-IDF v2 (Consumer-Grade Incentive Design Forensic engine) is an AI-powered audit "
    "tool that analyses a mobile or web application's user interface for incentive design "
    "patterns — mechanisms that drive engagement, monetisation, retention, or social "
    "behaviour. It also detects dark patterns: manipulative design tactics that exploit "
    "psychological biases against the user's own interests."
)
add_body(
    "An analyst uploads screenshots of key app surfaces through a Streamlit web interface. "
    "The engine automatically runs a five-stage AI pipeline and produces a fully structured, "
    "scored report that maps every detected pattern to the specific evidence that supports it."
)

add_callout(
    "Key value proposition: CG-IDF v2 replaces manual UX audits with a reproducible, "
    "evidence-linked, two-LLM verification pipeline — reducing both analyst time and the "
    "risk of unsubstantiated findings."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — APPLICATION STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("2. Application Structure", level=1)
add_divider()
add_body(
    "The Streamlit front-end exposes two primary screens: Evidence Collection and Audit Results. "
    "These screens act as a thin shell over the five-stage backend pipeline."
)

# 2.1
add_heading("2.1  Evidence Collection Screen", level=2)
add_body(
    "This is the entry point of every audit run. The analyst configures the LLM provider, "
    "uploads screenshots of the app surfaces to examine, and launches the pipeline."
)

add_table(
    ["UI Element", "Purpose"],
    [
        ("Sidebar — LLM Provider",       "Choose OpenAI or Anthropic as the analysis engine"),
        ("Sidebar — Model Override",     "Optionally pin a specific model (e.g. gpt-4o, claude-sonnet-4-6)"),
        ("Sidebar — API Key reminder",   "Surfaces the OPENAI_API_KEY / ANTHROPIC_API_KEY requirement"),
        ("Add New Evidence panel",       "Upload one or more screenshots; each is auto-assigned an ID (ev_001, ev_002, …)"),
        ("Captured Items table",         "Confirms uploaded files and their IDs before the audit starts"),
        ("Start Incentive Audit button", "Triggers the full five-stage pipeline (red CTA button)"),
        ("Audit Complete badge",         "Replaces the trigger button once the pipeline run finishes"),
    ]
)

add_callout(
    "Design note: Evidence IDs are deterministic (ev_001 … ev_N), making every pipeline run "
    "fully reproducible and comparable across audits of the same product."
)

# 2.2
add_heading("2.2  Audit Results Screen", level=2)
add_body("The results screen is divided into four visual zones stacked vertically.")

add_heading("2.2.1  Score Summary Bar", level=3)
add_body("Three top-level metrics give an immediate audit health signal:")
add_table(
    ["Metric", "Value (Example Run)", "Description"],
    [
        ("Overall Score",        "0.9",  "Mean of all five layer rollup scores (0.0 – 1.0)"),
        ("Flags Detected",       "5",    "Total number of pipeline issues raised"),
        ("Summary text",         "—",    "Auto-generated sentence: highest layer, flag breakdown"),
        ("Download JSON Report", "—",    "Full structured FinalReport as a downloadable file"),
    ]
)

add_heading("2.2.2  Incentive Profile — Radar Chart", level=3)
add_body(
    "A radar (spider) chart plots rollup scores for all five analysis layers simultaneously. "
    "The shape gives an immediate visual fingerprint of the app's incentive footprint: "
    "a layer that fills the outer ring (score ≈ 1.0) is saturated with detectable patterns; "
    "a collapsed axis means evidence was absent or inconclusive."
)

add_heading("2.2.3  Layer Scores — Bar Chart", level=3)
add_body("Horizontal bars give precise numeric rollup scores per layer:")
add_table(
    ["Layer", "Score (Example Run)"],
    [
        ("Engagement Incentives",       "0.93"),
        ("Monetisation Incentives",     "1.00"),
        ("Retention Incentives",        "1.00"),
        ("Social & Sharing Incentives", "1.00"),
        ("Dark Pattern Detection",      "0.65"),
    ]
)

add_heading("2.2.4  Audit Flags Panel", level=3)
add_body(
    "Collapsible flag groups list every issue raised during the pipeline. Each group shows "
    "the flag type and count; expanding it reveals the specific question IDs affected."
)
add_table(
    ["Flag Code", "Count", "Meaning"],
    [
        ("MISSING_SURFACE",  "1", "A required surface (home feed / checkout / onboarding) was not uploaded"),
        ("MISSING_ANSWER",   "2", "The LLM could not answer a question from the available evidence"),
        ("LOW_CONFIDENCE",   "2", "The LLM answered but confidence fell below the 0.6 threshold"),
    ]
)

add_heading("2.2.5  Detailed Breakdown Panel", level=3)
add_body(
    "A layer selector lets the analyst drill into individual questions within any layer. "
    "Each question shows a status icon (green check = answered & verified; red question mark "
    "= missing or low-confidence), the question ID and text, and an expandable detail "
    "section with the full LLM answer, confidence score, evidence references, and any "
    "verification notes added by Provider B."
)
add_body("Example questions visible in the Engagement layer:")
add_bullet("ENG_01: What UI patterns drive the user to spend more time in the app?")
add_bullet("ENG_02: Are there infinite scroll or auto-play mechanisms present?")
add_bullet("ENG_03: Does the app use variable-reward mechanics (e.g., likes, streaks)?")
add_bullet("ENG_04: Are notification or badge counts visible on the main surface?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — BACKEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3. Backend Architecture", level=1)
add_divider()

add_heading("3.1  Pipeline Overview", level=2)
add_body(
    "The engine is built on LangGraph — a state-machine orchestration framework. "
    "The pipeline is a directed acyclic graph (DAG) with five stages and one conditional branch:"
)

add_code_block(
    "Evidence + Screenshots\n"
    "       │\n"
    "       ▼\n"
    "  ┌─────────────────┐\n"
    "  │  Provider A     │   Multimodal LLM — answers 20 questions across 5 layers\n"
    "  └────────┬────────┘\n"
    "           │\n"
    "           ▼\n"
    "  ┌─────────────────┐\n"
    "  │  Rules Engine   │   Deterministic Python — flags issues, builds review queue\n"
    "  └────────┬────────┘\n"
    "           │\n"
    "    ┌──────┴──────┐\n"
    " [queue > 0]  [queue empty]\n"
    "       │              │\n"
    "       ▼              │\n"
    "  ┌─────────────────┐ │\n"
    "  │  Provider B     │ │   Text-only LLM — verifies flagged answers\n"
    "  └────────┬────────┘ │\n"
    "           │          │\n"
    "           └────┬─────┘\n"
    "                ▼\n"
    "  ┌─────────────────────┐\n"
    "  │  Merge + Scoring    │   Adjusts confidence, computes scores, builds FinalReport\n"
    "  └─────────────────────┘\n"
    "                ▼\n"
    "         FinalReport (JSON)"
)

add_heading("3.2  File Map", level=2)
add_code_block(
    "cg-idf-v2/\n"
    "├── app.py                  # Streamlit UI\n"
    "├── graph.py                # LangGraph DAG + conditional routing\n"
    "├── schema.py               # Pydantic models (AuditState, FinalReport, Layer, …)\n"
    "├── llm.py                  # Unified OpenAI / Anthropic client\n"
    "├── main.py                 # CLI entry point\n"
    "└── nodes/\n"
    "    ├── provider_a.py       # Step 2 — multimodal LLM analysis\n"
    "    ├── rules_engine.py     # Step 3 — deterministic validation\n"
    "    ├── provider_b.py       # Step 4 — text-only verification\n"
    "    └── merge_scoring.py    # Step 5 — confidence adjustment + scoring"
)

add_heading("3.3  Node-by-Node Summary", level=2)

add_heading("Provider A — Primary LLM Analysis", level=3)
add_body(
    "Provider A is the first and most expensive LLM call. It receives all uploaded evidence "
    "items together with their base64-encoded screenshots and answers 20 structured questions "
    "spread across five analysis layers (4 questions per layer). Every answer is classified as:"
)
add_bullet("supported", bold_prefix="supported  — ")
add_bullet("directly backed by at least one cited evidence reference")
add_bullet("inferred", bold_prefix="inferred  — ")
add_bullet("reasonable deduction without a direct screenshot reference")
add_bullet("unknown", bold_prefix="unknown  — ")
add_bullet("insufficient evidence to form a conclusion")
add_body(
    "The node automatically demotes a 'supported' claim to 'inferred' if no evidence_refs "
    "were actually cited, preventing the LLM from over-claiming evidence quality.",
    space_after=10
)

add_heading("Rules Engine — Deterministic Gating", level=3)
add_body(
    "The rules engine is pure Python — no LLM involved. It applies five deterministic checks "
    "to Provider A's output and is the primary quality gate of the pipeline:"
)
add_table(
    ["Rule", "Flag Raised", "Threshold"],
    [
        ("Required surface coverage",  "MISSING_SURFACE",     "home_feed, checkout, onboarding must be present"),
        ("Unsupported claim check",    "UNSUPPORTED_CLAIM",   "answered 'supported' but no evidence_refs cited"),
        ("Missing answer check",       "MISSING_ANSWER",      "no answer_text or answer_type = unknown"),
        ("Low confidence check",       "LOW_CONFIDENCE",      "confidence < 0.6"),
        ("Incomplete layer coverage",  "LAYER_MISSING",       "any of the 5 required layers absent"),
    ]
)
add_body(
    "Items that trigger a flag are added to the review_queue. The graph's conditional router "
    "sends the pipeline to Provider B only if this queue is non-empty, keeping costs low "
    "when the first-pass analysis is clean."
)

add_heading("Provider B — Verification LLM", level=3)
add_body(
    "Provider B is a targeted, text-only verification step. It receives only the flagged "
    "items from the review queue — not the full set of 20 questions — and returns one of "
    "five verdicts for each:"
)
add_table(
    ["Verdict", "Meaning"],
    [
        ("confirm",               "Answer well-supported; no change needed"),
        ("downgrade",             "Partially correct; confidence should be reduced"),
        ("contradiction",         "Evidence contradicts the answer"),
        ("insufficient_evidence", "Evidence too thin to support the claim"),
        ("missing_evidence",      "Referenced evidence absent or uninformative"),
    ]
)

add_heading("Merge + Scoring — Final Assembly", level=3)
add_body(
    "The merge node applies Provider B verdicts to the question objects and then computes "
    "the scoring hierarchy:"
)
add_table(
    ["Verdict Applied", "Confidence Effect", "Side Effect"],
    [
        ("confirm",               "No change",                    "—"),
        ("downgrade",             "× 0.6 (floor: 0.1)",          "Note appended to question"),
        ("contradiction",         "Forced to 0.1",               "AuditFlag raised + note"),
        ("insufficient_evidence", "× 0.6 (floor: 0.1)",          "Note appended to question"),
        ("missing_evidence",      "× 0.6 (floor: 0.1)",          "Note appended to question"),
    ]
)
add_body(
    "Layer rollup = mean confidence across all questions in the layer. Unanswered or "
    "'unknown' questions contribute 0.0, so gaps actively penalise the score rather than "
    "being silently excluded. Overall score = mean of all five layer rollups."
)

add_heading("3.4  Data Model Flow", level=2)
add_code_block(
    "Evidence[]  ──▶  AuditState\n"
    "                    ├── screen_facts[]          (extracted by Provider A)\n"
    "                    ├── layers{}                (5 layers × 4 questions each)\n"
    "                    │       └── questions[]\n"
    "                    │               ├── llm_answer\n"
    "                    │               ├── answer_type\n"
    "                    │               ├── confidence\n"
    "                    │               └── evidence_refs[]\n"
    "                    ├── review_queue[]           (built by Rules Engine)\n"
    "                    ├── verifications[]          (returned by Provider B)\n"
    "                    ├── pipeline_flags[]         (rules + contradiction flags)\n"
    "                    └── final_report             (assembled by Merge node)\n"
    "                            ├── layers{}  (with rollup_scores)\n"
    "                            ├── flags[]\n"
    "                            ├── overall_score\n"
    "                            └── summary (auto-generated text)"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — OUTPUT ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

add_heading("4. Output Analysis — Example Run", level=1)
add_divider()
add_body(
    "The example audit processed a single screenshot (ev_001) and produced scores across "
    "all five layers. The table below interprets the key findings:"
)

add_table(
    ["Observation", "Interpretation"],
    [
        ("Monetisation, Retention, Social all at 1.00",
         "Provider A found strong, high-confidence signals in the single screenshot"),
        ("Dark Pattern Detection at 0.65",
         "Partial evidence — some patterns were inferred rather than directly supported"),
        ("1× MISSING_SURFACE",
         "The screenshot did not cover one of the three required surfaces (home feed, checkout, or onboarding)"),
        ("2× MISSING_ANSWER",
         "Two questions had no answerable evidence in the uploaded material"),
        ("2× LOW_CONFIDENCE",
         "Two answers fell below the 0.6 confidence threshold and were sent to Provider B"),
        ("Overall score 0.9",
         "High density of detectable incentive patterns across the audited surface"),
        ("Highest layer: 'Monetisation Incentives'",
         "The screenshot contained the strongest monetisation signals of any layer"),
    ]
)

add_callout(
    "Important — score direction: A high score does NOT mean 'good design.' It means the "
    "app is dense with detectable incentive mechanisms. A score of 1.00 on Dark Pattern "
    "Detection is a serious red flag, not a positive result."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — WHAT CAN BE IMPROVED IN THE OUTPUT
# ══════════════════════════════════════════════════════════════════════════════

add_heading("5. What Can Be Improved in the Output", level=1)
add_divider()
add_body(
    "The pipeline logic and data model are sound. The main improvement opportunities are "
    "in how the output is communicated to human reviewers."
)

improvements = [
    (
        "1.  Score Interpretation Guidance",
        BLUE,
        [
            "Current state: Numeric scores (0.93, 1.00, 0.65, …) are displayed without contextual meaning.",
            "Improvement: Add threshold bands with colour coding:",
            "    🟢  0.0 – 0.3  Low pattern density",
            "    🟡  0.3 – 0.6  Moderate density",
            "    🔴  0.6 – 1.0  High density — warrants review",
            "Critical for Dark Pattern Detection specifically: a high score there should surface "
            "a warning callout rather than rendering identically to a high Engagement score.",
        ]
    ),
    (
        "2.  Human-Readable Flag Labels",
        BLUE,
        [
            "Current state: Flags display as FlagCode.MISSING_ANSWER (2) — raw Python enum names.",
            "Improvement: Strip the FlagCode. prefix and add plain-English descriptions:",
            "    Missing Answer (2) — The LLM could not find evidence to answer these questions.",
            "    Low Confidence (2) — These answers may be unreliable; consider uploading more screenshots.",
        ]
    ),
    (
        "3.  Inline Evidence Thumbnails",
        BLUE,
        [
            "Current state: The Detailed Breakdown shows answers and evidence IDs as text only.",
            "Improvement: Display a thumbnail of the referenced screenshot next to each question "
            "answer so a human reviewer can immediately verify the claim visually without switching "
            "context to a file browser.",
        ]
    ),
    (
        "4.  Dedicated Contradictions Panel",
        BLUE,
        [
            "Current state: Contradictions are present in the JSON output but have no dedicated UI section.",
            "Improvement: Add a 'Contradictions Detected' panel that appears only when contradictions[] "
            "is non-empty, showing Provider A's original answer and Provider B's counter-finding "
            "side by side for easy human arbitration.",
        ]
    ),
    (
        "5.  Limited-Evidence Warning Banner",
        BLUE,
        [
            "Current state: Three layers scored 1.00 from a single screenshot — statistically fragile.",
            "Improvement: Show a dismissible banner when fewer than 3 evidence items are provided:",
            "    'Audit based on limited evidence. Scores may not reflect the full product.'",
            "Consider weighting rollup scores by evidence count to dampen over-confident results.",
        ]
    ),
    (
        "6.  Additional Export Formats",
        BLUE,
        [
            "Current state: Only a JSON download is available.",
            "Improvement:",
            "    PDF export — formatted report with the radar chart embedded, suitable for sharing.",
            "    CSV export — question-level data for stakeholders who prefer spreadsheet analysis.",
        ]
    ),
    (
        "7.  Inverted Radar Axis for Dark Patterns",
        BLUE,
        [
            "Current state: All five radar axes point outward equally, making a 'bigger shape' "
            "appear uniformly positive.",
            "Improvement: Invert the Dark Pattern Detection axis (or render it in red) so the "
            "radar chart intuitively represents a healthy product: large engagement / retention "
            "footprint, small dark-pattern footprint.",
        ]
    ),
    (
        "8.  Audit History & Version Comparison",
        BLUE,
        [
            "Current state: Each run is standalone with no persistence between sessions.",
            "Improvement: Store run history in a local SQLite or file-based store and allow "
            "overlaying two radar charts — enabling before/after comparisons after a product "
            "update, or competitor benchmarking.",
        ]
    ),
]

for title, color, bullets in improvements:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    run.font.bold  = True
    run.font.size  = Pt(12)
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    for b in bullets:
        add_bullet(b)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. Summary", level=1)
add_divider()
add_body(
    "CG-IDF v2 is a well-structured, evidence-linked incentive audit system. Its core "
    "strengths are the deterministic rules engine (which prevents LLM hallucinations from "
    "silently inflating scores), the conditional Provider B verification step (cost-efficient "
    "and targeted), and the fully structured JSON output (machine-readable and suitable for "
    "downstream integrations)."
)
add_body(
    "The areas most worth improving are in output communication — making scores interpretable, "
    "flags human-readable, and evidence visually traceable — rather than in the underlying "
    "pipeline logic, which is architecturally sound."
)

add_table(
    ["Strength", "Improvement Opportunity"],
    [
        ("Two-LLM review with deterministic gating",  "Score threshold bands + direction for Dark Patterns"),
        ("Evidence-linked answers (supported / inferred / unknown)", "Inline screenshot thumbnails per answer"),
        ("Conditional routing (cost-efficient)",       "Limited-evidence warning banner"),
        ("Structured FinalReport JSON",                "PDF and CSV export options"),
        ("Reproducible evidence IDs",                  "Audit history + radar overlay comparison"),
    ]
)

add_callout(
    "Pipeline Run ID referenced in example output: 0168fc35-52f1-4257-bb93-1f11b6769772\n"
    "Report generated from CG-IDF v2 source code and UI screenshots — March 2026."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════

out = "/home/user/Atlas_DS/cg-idf-v2/CG-IDF-v2-Report.docx"
doc.save(out)
print(f"Saved: {out}")
